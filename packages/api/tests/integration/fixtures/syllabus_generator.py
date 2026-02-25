"""Programmatic test syllabus generators.

Each function returns (file_bytes, mime_type, ground_truth_dict) where
ground_truth contains the expected extraction fields for accuracy measurement.
"""

from __future__ import annotations

# ---------------------------------------------------------------------------
# Minimal PDF builder (extended from conftest.py pattern)
# ---------------------------------------------------------------------------


def _text_to_pdf(lines: list[str]) -> bytes:
    """Convert lines of text to a valid single-page PDF."""
    ops = ["BT", "/F1 10 Tf", "72 750 Td", "12 TL"]
    for line in lines:
        escaped = (
            line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        )
        ops.append(f"({escaped}) '")
    ops.append("ET")
    stream = "\n".join(ops).encode("latin-1")

    obj1 = b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
    obj2 = b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
    obj3 = (
        b"3 0 obj\n"
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]\n"
        b"   /Contents 4 0 R\n"
        b"   /Resources << /Font << /F1 5 0 R >> >> >>\n"
        b"endobj\n"
    )
    obj4 = (
        f"4 0 obj\n<< /Length {len(stream)} >>\nstream\n".encode()
        + stream
        + b"\nendstream\nendobj\n"
    )
    obj5 = (
        b"5 0 obj\n"
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\n"
        b"endobj\n"
    )
    objects = [obj1, obj2, obj3, obj4, obj5]
    header = b"%PDF-1.4\n"
    offsets: list[int] = []
    pos = len(header)
    for obj in objects:
        offsets.append(pos)
        pos += len(obj)
    xref_start = pos
    xref = b"xref\n"
    xref += f"0 {len(objects) + 1}\n".encode()
    xref += b"0000000000 65535 f \n"
    for offset in offsets:
        xref += f"{offset:010d} 00000 n \n".encode()
    trailer = (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_start}\n%%EOF\n"
    ).encode()
    return header + b"".join(objects) + xref + trailer


# ---------------------------------------------------------------------------
# Syllabus 1: Clean STEM syllabus (baseline)
# ---------------------------------------------------------------------------

STEM_SEMESTER = {
    "start": "2026-01-12",
    "end": "2026-05-01",
    "meeting_days": ["tuesday", "thursday"],
}


def build_stem_syllabus() -> tuple[bytes, str, dict]:
    """PHYS 201 - University Physics I. Clean, well-structured STEM syllabus."""
    lines = [
        "PHYS 201: University Physics I - Thermodynamics",
        "Spring 2026 Syllabus",
        "",
        "Instructor: Dr. Richard Feynman",
        "Email: rfeynman@caltech.edu",
        "Office Hours: Tuesday/Thursday 2:00-3:30 PM, Room 214 Sloane",
        "",
        "Semester: January 12 - May 1, 2026",
        "Meeting Times: Tuesday and Thursday, 10:00 - 11:30 AM",
        "Location: Room 201, Downs Laboratory",
        "",
        "Grade Breakdown:",
        "  Midterm Exam 1: 15%",
        "  Midterm Exam 2: 15%",
        "  Final Exam: 30%",
        "  Laboratory Reports: 20%",
        "  Homework Assignments: 15% (lowest dropped)",
        "  Participation: 5%",
        "",
        "Course Schedule:",
        "  Week 1 (Jan 12-16): Temperature and Heat",
        "  Week 2 (Jan 19-23): Ideal Gas Law",
        "  Week 3 (Jan 26-30): Kinetic Theory",
        "  Week 4 (Feb 2-6): First Law of Thermodynamics",
        "  Week 5 (Feb 9-13): Heat Engines and Refrigerators",
        "  Week 6 (Feb 16-20): Midterm 1 Review + Exam",
        "  Week 7 (Feb 23-27): Second Law of Thermodynamics",
        "  Week 8 (Mar 2-6): Entropy",
        "  Week 9 (Mar 9-13): Spring Break - No Class",
        "  Week 10 (Mar 16-20): Thermodynamic Potentials",
        "  Week 11 (Mar 23-27): Phase Transitions",
        "  Week 12 (Mar 30-Apr 3): Midterm 2 Review + Exam",
        "  Week 13 (Apr 6-10): Statistical Mechanics Intro",
        "  Week 14 (Apr 13-17): Boltzmann Distribution",
        "  Week 15 (Apr 20-24): Quantum Statistics",
        "  Week 16 (Apr 27-May 1): Review for Final",
        "",
        "Assessment Details:",
        "  Midterm Exam 1 - February 19, 2026 - Chapters 1-5 - 15%",
        "  Midterm Exam 2 - April 2, 2026 - Chapters 6-10 - 15%",
        "  Final Exam - May 7, 2026 - Cumulative - 30%",
        "  Lab Report 1 - February 5, 2026 - Calorimetry - 5%",
        "  Lab Report 2 - March 5, 2026 - Heat Engines - 5%",
        "  Lab Report 3 - April 16, 2026 - Phase Transitions - 5%",
        "  Lab Report 4 - April 30, 2026 - Statistical Mechanics - 5%",
        "  Homework 1 - January 22, 2026 - Ch. 1 - 2.5%",
        "  Homework 2 - February 5, 2026 - Ch. 2-3 - 2.5%",
        "  Homework 3 - February 26, 2026 - Ch. 4-5 - 2.5%",
        "  Homework 4 - March 19, 2026 - Ch. 6-7 - 2.5%",
        "  Homework 5 - April 9, 2026 - Ch. 8-9 - 2.5%",
        "  Homework 6 - April 23, 2026 - Ch. 10-11 - 2.5%",
        "",
        "Late Policy: 10% penalty per day, maximum 3 days late.",
        "Academic Integrity: Violations result in course failure.",
        "Attendance: Mandatory. Three unexcused absences lower final grade by one step.",
    ]

    ground_truth = {
        "course_name": "University Physics I",
        "course_code": "PHYS 201",
        "instructor_name": "Dr. Richard Feynman",
        "instructor_email": "rfeynman@caltech.edu",
        "assessment_count": 13,
        "assessments": [
            {
                "title": "Midterm Exam 1",
                "type": "exam",
                "due_date": "2026-02-19",
                "weight_percent": 15.0,
            },
            {
                "title": "Midterm Exam 2",
                "type": "exam",
                "due_date": "2026-04-02",
                "weight_percent": 15.0,
            },
            {
                "title": "Final Exam",
                "type": "exam",
                "due_date": "2026-05-07",
                "weight_percent": 30.0,
            },
            {
                "title": "Lab Report 1",
                "type": "lab",
                "due_date": "2026-02-05",
                "weight_percent": 5.0,
            },
            {
                "title": "Lab Report 2",
                "type": "lab",
                "due_date": "2026-03-05",
                "weight_percent": 5.0,
            },
            {
                "title": "Lab Report 3",
                "type": "lab",
                "due_date": "2026-04-16",
                "weight_percent": 5.0,
            },
            {
                "title": "Lab Report 4",
                "type": "lab",
                "due_date": "2026-04-30",
                "weight_percent": 5.0,
            },
            {
                "title": "Homework 1",
                "type": "homework",
                "due_date": "2026-01-22",
                "weight_percent": 2.5,
            },
            {
                "title": "Homework 2",
                "type": "homework",
                "due_date": "2026-02-05",
                "weight_percent": 2.5,
            },
            {
                "title": "Homework 3",
                "type": "homework",
                "due_date": "2026-02-26",
                "weight_percent": 2.5,
            },
            {
                "title": "Homework 4",
                "type": "homework",
                "due_date": "2026-03-19",
                "weight_percent": 2.5,
            },
            {
                "title": "Homework 5",
                "type": "homework",
                "due_date": "2026-04-09",
                "weight_percent": 2.5,
            },
            {
                "title": "Homework 6",
                "type": "homework",
                "due_date": "2026-04-23",
                "weight_percent": 2.5,
            },
        ],
        "grade_breakdown_total": 100.0,
        "date_count": 13,
        "semester": STEM_SEMESTER,
    }

    return _text_to_pdf(lines), "application/pdf", ground_truth


# ---------------------------------------------------------------------------
# Syllabus 2: Business school with tables
# ---------------------------------------------------------------------------

BUSINESS_SEMESTER = {
    "start": "2026-01-12",
    "end": "2026-05-01",
    "meeting_days": ["monday", "wednesday"],
}


def build_business_syllabus() -> tuple[bytes, str, dict]:
    """BUS 301 - Strategic Management. Tables for grade breakdown, mixed dates."""
    lines = [
        "BUS 301: Strategic Management",
        "Spring 2026 Course Syllabus",
        "",
        "Professor: Dr. Sarah Chen",
        "Email: schen@wharton.edu",
        "Office: Room 512 Huntsman Hall",
        "Office Hours: Monday 3-5 PM",
        "",
        "Semester: January 12 - May 1, 2026",
        "Class: Mon/Wed 1:00 - 2:30 PM",
        "",
        "| Component            | Weight | Notes                    |",
        "|----------------------|--------|--------------------------|",
        "| Case Analysis (x4)   | 20%    | Individual write-ups     |",
        "| Midterm Exam         | 25%    | In-class, March 4        |",
        "| Group Project        | 30%    | Final presentation       |",
        "| Class Participation  | 15%    | Attendance + discussion  |",
        "| Final Exam           | 10%    | Take-home, due May 1     |",
        "",
        "Assessment Schedule:",
        "  Case Analysis 1 - January 28, 2026 - Industry Analysis - 5%",
        "  Case Analysis 2 - February 18, 2026 - Competitive Advantage - 5%",
        "  Midterm Exam - March 4, 2026 - Chapters 1-6 - 25%",
        "  Case Analysis 3 - Week 10 Wednesday - Corporate Strategy - 5%",
        "  Case Analysis 4 - April 8, 2026 - International Strategy - 5%",
        "  Group Project Proposal - February 25, 2026 - 5%",
        "  Group Project Presentation - April 27, 2026 - 25%",
        "  Final Exam - May 1, 2026 - Comprehensive - 10%",
        "",
        "Late Policy: Case analyses accepted up to 24 hours late for 20% penalty.",
        "Academic Integrity: All work must comply with the Honor Code.",
    ]

    ground_truth = {
        "course_name": "Strategic Management",
        "course_code": "BUS 301",
        "instructor_name": "Dr. Sarah Chen",
        "instructor_email": "schen@wharton.edu",
        "assessment_count": 8,
        "assessments": [
            {
                "title": "Case Analysis 1",
                "type": "homework",
                "due_date": "2026-01-28",
                "weight_percent": 5.0,
            },
            {
                "title": "Case Analysis 2",
                "type": "homework",
                "due_date": "2026-02-18",
                "weight_percent": 5.0,
            },
            {
                "title": "Midterm Exam",
                "type": "exam",
                "due_date": "2026-03-04",
                "weight_percent": 25.0,
            },
            {
                "title": "Case Analysis 3",
                "type": "homework",
                "due_date": "2026-03-18",  # Week 10 Wednesday
                "weight_percent": 5.0,
            },
            {
                "title": "Case Analysis 4",
                "type": "homework",
                "due_date": "2026-04-08",
                "weight_percent": 5.0,
            },
            {
                "title": "Group Project Proposal",
                "type": "project",
                "due_date": "2026-02-25",
                "weight_percent": 5.0,
            },
            {
                "title": "Group Project Presentation",
                "type": "presentation",
                "due_date": "2026-04-27",
                "weight_percent": 25.0,
            },
            {
                "title": "Final Exam",
                "type": "exam",
                "due_date": "2026-05-01",
                "weight_percent": 10.0,
            },
        ],
        "grade_breakdown_total": 100.0,
        "date_count": 8,
        "semester": BUSINESS_SEMESTER,
    }

    return _text_to_pdf(lines), "application/pdf", ground_truth


# ---------------------------------------------------------------------------
# Syllabus 3: Minimal 2-page syllabus
# ---------------------------------------------------------------------------

MINIMAL_SEMESTER = {
    "start": "2026-01-12",
    "end": "2026-05-01",
    "meeting_days": ["tuesday", "thursday"],
}


def build_minimal_syllabus() -> tuple[bytes, str, dict]:
    """ENG 102 - English Composition. Sparse: 3 assessments, no schedule."""
    lines = [
        "ENG 102: English Composition II",
        "Spring 2026",
        "",
        "Instructor: Prof. Maria Rodriguez",
        "",
        "Assignments and Grading:",
        "  Essay 1 (Argumentative) - March 6, 2026 - 40%",
        "  Essay 2 (Research Paper) - April 24, 2026 - 40%",
        "  Class Participation - Ongoing - 20%",
        "",
        "All essays must be submitted via the online portal by 11:59 PM.",
    ]

    ground_truth = {
        "course_name": "English Composition II",
        "course_code": "ENG 102",
        "instructor_name": "Prof. Maria Rodriguez",
        "assessment_count": 3,
        "assessments": [
            {
                "title": "Essay 1",
                "type": "paper",
                "due_date": "2026-03-06",
                "weight_percent": 40.0,
            },
            {
                "title": "Essay 2",
                "type": "paper",
                "due_date": "2026-04-24",
                "weight_percent": 40.0,
            },
            {
                "title": "Class Participation",
                "type": "participation",
                "due_date": None,
                "weight_percent": 20.0,
            },
        ],
        "grade_breakdown_total": 100.0,
        "date_count": 2,
        "semester": MINIMAL_SEMESTER,
    }

    return _text_to_pdf(lines), "application/pdf", ground_truth


# ---------------------------------------------------------------------------
# Syllabus 4: "Week N" date format throughout
# ---------------------------------------------------------------------------

WEEK_FORMAT_SEMESTER = {
    "start": "2026-01-12",
    "end": "2026-05-01",
    "meeting_days": ["tuesday", "thursday"],
}


def build_week_format_syllabus() -> tuple[bytes, str, dict]:
    """CS 340 - Algorithms. All dates use 'Week N DayName' format."""
    lines = [
        "CS 340: Design and Analysis of Algorithms",
        "Spring 2026",
        "",
        "Instructor: Dr. Alan Turing",
        "Email: aturing@mit.edu",
        "Office Hours: Thursday 3-5 PM, Gates 432",
        "",
        "Semester: January 12 - May 1, 2026",
        "Lectures: Tuesday and Thursday 9:00 - 10:30 AM",
        "",
        "Grade Breakdown:",
        "  Problem Sets (6): 30% (lowest dropped)",
        "  Midterm Exam: 25%",
        "  Final Exam: 35%",
        "  Participation: 10%",
        "",
        "Assessment Schedule:",
        "  Problem Set 1 - Week 3 Tuesday - Divide and Conquer - 6%",
        "  Problem Set 2 - Week 5 Thursday - Graph Algorithms - 6%",
        "  Problem Set 3 - Week 7 Tuesday - Dynamic Programming - 6%",
        "  Midterm Exam - Week 8 Thursday - All topics through DP - 25%",
        "  Problem Set 4 - Week 10 Tuesday - Network Flow - 6%",
        "  Problem Set 5 - Week 12 Thursday - NP-Completeness - 6%",
        "  Problem Set 6 - Week 14 Tuesday - Approximation Algorithms - 6%",
        "  Final Exam - Week 16 Thursday - Cumulative - 35%",
        "",
        "Course Schedule:",
        "  Week 1: Introduction, Asymptotic Analysis",
        "  Week 2: Recurrences, Master Theorem",
        "  Week 3: Divide and Conquer (Mergesort, Quicksort)",
        "  Week 4: Sorting Lower Bounds, Linear-time Sorting",
        "  Week 5: Graph Basics, BFS, DFS",
        "  Week 6: Shortest Paths, Minimum Spanning Trees",
        "  Week 7: Dynamic Programming I",
        "  Week 8: Dynamic Programming II + Midterm",
        "  Week 9: Spring Break",
        "  Week 10: Network Flow",
        "  Week 11: Bipartite Matching",
        "  Week 12: NP-Completeness",
        "  Week 13: Reductions",
        "  Week 14: Approximation Algorithms",
        "  Week 15: Randomized Algorithms",
        "  Week 16: Review + Final Exam",
        "",
        "Late Policy: Problem sets due at 11:59 PM. 25% off per day, max 2 days.",
    ]

    # Week 1 starts Jan 12 (Monday). Tue = Jan 13, Thu = Jan 15.
    # Week N Tue = Jan 13 + (N-1)*7, Week N Thu = Jan 15 + (N-1)*7
    ground_truth = {
        "course_name": "Design and Analysis of Algorithms",
        "course_code": "CS 340",
        "instructor_name": "Dr. Alan Turing",
        "instructor_email": "aturing@mit.edu",
        "assessment_count": 8,
        "assessments": [
            {
                "title": "Problem Set 1",
                "type": "homework",
                "due_date": "2026-01-27",  # Week 3 Tue
                "weight_percent": 6.0,
            },
            {
                "title": "Problem Set 2",
                "type": "homework",
                "due_date": "2026-02-12",  # Week 5 Thu
                "weight_percent": 6.0,
            },
            {
                "title": "Problem Set 3",
                "type": "homework",
                "due_date": "2026-02-24",  # Week 7 Tue
                "weight_percent": 6.0,
            },
            {
                "title": "Midterm Exam",
                "type": "exam",
                "due_date": "2026-03-05",  # Week 8 Thu
                "weight_percent": 25.0,
            },
            {
                "title": "Problem Set 4",
                "type": "homework",
                "due_date": "2026-03-17",  # Week 10 Tue
                "weight_percent": 6.0,
            },
            {
                "title": "Problem Set 5",
                "type": "homework",
                "due_date": "2026-04-02",  # Week 12 Thu
                "weight_percent": 6.0,
            },
            {
                "title": "Problem Set 6",
                "type": "homework",
                "due_date": "2026-04-14",  # Week 14 Tue
                "weight_percent": 6.0,
            },
            {
                "title": "Final Exam",
                "type": "exam",
                "due_date": "2026-04-30",  # Week 16 Thu
                "weight_percent": 35.0,
            },
        ],
        "grade_breakdown_total": 100.0,
        "date_count": 8,
        "semester": WEEK_FORMAT_SEMESTER,
    }

    return _text_to_pdf(lines), "application/pdf", ground_truth


# ---------------------------------------------------------------------------
# Syllabus 5: DOCX file
# ---------------------------------------------------------------------------

DOCX_SEMESTER = {
    "start": "2026-01-12",
    "end": "2026-05-01",
    "meeting_days": ["monday", "wednesday", "friday"],
}


def build_docx_syllabus() -> tuple[bytes, str, dict]:
    """HIST 215 - Modern World History. Word document with headings + table."""
    from io import BytesIO

    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    doc.add_heading("HIST 215: Modern World History", level=1)
    doc.add_heading("Spring 2026 Syllabus", level=2)

    doc.add_paragraph("")
    doc.add_paragraph("Instructor: Dr. Howard Zinn")
    doc.add_paragraph("Email: hzinn@bu.edu")
    doc.add_paragraph("Office Hours: Wednesday 1-3 PM, Room 226 Bay State")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Semester: January 12 - May 1, 2026. "
        "Class meets Monday, Wednesday, Friday 11:00 AM - 12:00 PM."
    )

    doc.add_heading("Grade Breakdown", level=2)
    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"
    headers = ["Component", "Weight", "Details"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ("Reading Responses (10)", "20%", "Weekly, 2 pages each"),
        ("Midterm Essay", "20%", "In-class, March 6"),
        ("Research Paper", "30%", "Due April 24"),
        ("Final Exam", "20%", "May 8, comprehensive"),
        ("Participation", "10%", "Discussion and attendance"),
    ]
    for row_idx, (comp, weight, details) in enumerate(data, start=1):
        table.rows[row_idx].cells[0].text = comp
        table.rows[row_idx].cells[1].text = weight
        table.rows[row_idx].cells[2].text = details

    doc.add_heading("Assessment Details", level=2)
    doc.add_paragraph(
        "Midterm Essay - March 6, 2026 - "
        "Compare two revolutions from the 18th-19th century - 20%"
    )
    doc.add_paragraph(
        "Research Paper - April 24, 2026 - "
        "Original research on a 20th century topic (15-20 pages) - 30%"
    )
    doc.add_paragraph(
        "Final Exam - May 8, 2026 - "
        "Comprehensive essay exam covering all course material - 20%"
    )

    doc.add_heading("Course Schedule", level=2)
    schedule_items = [
        "Week 1 (Jan 12-16): Introduction - What is Modern History?",
        "Week 2 (Jan 19-23): The Enlightenment",
        "Week 3 (Jan 26-30): American and French Revolutions",
        "Week 4 (Feb 2-6): Industrial Revolution",
        "Week 5 (Feb 9-13): Imperialism and Colonialism",
        "Week 6 (Feb 16-20): World War I",
        "Week 7 (Feb 23-27): Russian Revolution",
        "Week 8 (Mar 2-6): Interwar Period + Midterm",
        "Week 9 (Mar 9-13): Spring Break",
        "Week 10 (Mar 16-20): World War II",
        "Week 11 (Mar 23-27): The Cold War",
        "Week 12 (Mar 30-Apr 3): Decolonization",
        "Week 13 (Apr 6-10): Civil Rights Movements",
        "Week 14 (Apr 13-17): Globalization",
        "Week 15 (Apr 20-24): The Digital Age",
        "Week 16 (Apr 27-May 1): Review",
    ]
    for item in schedule_items:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("Policies", level=2)
    doc.add_paragraph(
        "Late Policy: Papers lose one letter grade per day late. "
        "Extensions must be requested 48 hours in advance."
    )
    doc.add_paragraph(
        "Academic Integrity: Plagiarism results in an automatic F for the assignment."
    )

    buf = BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    mime = (
        "application/vnd.openxmlformats-officedocument"
        ".wordprocessingml.document"
    )

    ground_truth = {
        "course_name": "Modern World History",
        "course_code": "HIST 215",
        "instructor_name": "Dr. Howard Zinn",
        "instructor_email": "hzinn@bu.edu",
        "assessment_count": 3,
        "assessments": [
            {
                "title": "Midterm Essay",
                "type": "exam",
                "due_date": "2026-03-06",
                "weight_percent": 20.0,
            },
            {
                "title": "Research Paper",
                "type": "paper",
                "due_date": "2026-04-24",
                "weight_percent": 30.0,
            },
            {
                "title": "Final Exam",
                "type": "exam",
                "due_date": "2026-05-08",
                "weight_percent": 20.0,
            },
        ],
        "grade_breakdown_total": 100.0,
        "date_count": 3,
        "semester": DOCX_SEMESTER,
    }

    return docx_bytes, mime, ground_truth


# ---------------------------------------------------------------------------
# Convenience: all syllabi
# ---------------------------------------------------------------------------

ALL_SYLLABI = {
    "stem": build_stem_syllabus,
    "business": build_business_syllabus,
    "minimal": build_minimal_syllabus,
    "week_format": build_week_format_syllabus,
    "docx": build_docx_syllabus,
}
